import docx

firstset = [151,152,1,2,153,3,4,5,154,155,6,156,7,157,8,158,9,159,10,160,11,12,161,162,13,163,14,164,165,15,166,167,16,168,169,17,170,18,171,172,173,174,19,20,175,21,176,177,22,178]
secondset = [23,24,179,180,25,181,26,27,28,182,183,184,185,186,29,187,188,189,30,31,190,191,32,33,192,193,194,195,196,197,198,34,199,35,36,200,37,201,38,39,40,202,203,41,204,205,42,43,206,207]
thirdset = [44,208,209,45,210,211,46,212,213,47,214,48,49,50,215,216,51,217,52,53,218,219,54,220,55,56,57,221,58,222,223,59,60,61,62,63,64,65,224,66,225,67,68,69,226,70,227,228,71,72]
fourthset = [229,230,73,231,232,233,74,234,235,75,76,236,77,237,238,239,240,241,242,78,243,79,244,80,245,246,81,82,83,84,85,247,86,248,87,249,250,251,88,252,253,89,90,91,92,254,255,93,94,256,]
fifthset = [95,96,257,97,98,258,259,99,260,100,101,261,262,263,102,264,103,104,265,266,267,268,269,105,106,107,270,271,272,273,274,108,109,110,111,112,275,113,276,114,115,277,278,116,279,280,281,117,118,119]
sixthset = [120,121,122,123,282,283,124,125,126,127,128,129,284,130,285,131,132,133,134,135,286,136,287,137,288,138,289,139,140,290,291,141,142,143,292,293,144,294,295,145,296,146,147,148,297,149,150,298,299,300]


doc = docx.Document('nmdata1.docx')
for i in firstset:
	doc.add_picture('Figure' + str(i) + 'nm.jpg',width=docx.shared.Inches(16.21/2),height=docx.shared.Inches(12.15/2))
	doc.save('nmdata1.docx')
	

doc = docx.Document('nmdata2.docx')
for i in secondset:
	doc.add_picture('Figure' + str(i) + 'nm.jpg',width=docx.shared.Inches(16.21/2),height=docx.shared.Inches(12.15/2))
	doc.save('nmdata2.docx')
	
doc = docx.Document('nmdata3.docx')
for i in thirdset:
	doc.add_picture('Figure' + str(i) + 'nm.jpg',width=docx.shared.Inches(16.21/2),height=docx.shared.Inches(12.15/2))
	doc.save('nmdata3.docx')
	
doc = docx.Document('nmdata4.docx')
for i in fourthset:
	doc.add_picture('Figure' + str(i) + 'nm.jpg',width=docx.shared.Inches(16.21/2),height=docx.shared.Inches(12.15/2))
	doc.save('nmdata4.docx')
	
doc = docx.Document('nmdata5.docx')
for i in fifthset:
	doc.add_picture('Figure' + str(i) + 'nm.jpg',width=docx.shared.Inches(16.21/2),height=docx.shared.Inches(12.15/2))
	doc.save('nmdata5.docx')
	
doc = docx.Document('nmdata6.docx')
for i in sixthset:
	doc.add_picture('Figure' + str(i) + 'nm.jpg',width=docx.shared.Inches(16.21/2),height=docx.shared.Inches(12.15/2))
	doc.save('nmdata6.docx')
